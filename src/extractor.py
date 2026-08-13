from __future__ import annotations

from collections.abc import Iterator
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .config import DEFAULT_CONFIG, AppConfig, ExtractorConfig
from .models import ExtractedDocument, PageContent
from .utils import get_logger, repair_pdf_artifacts

logger = get_logger(__name__)


class ExtractionError(Exception):
    pass


def iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    seen: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id in seen:
                continue
            seen.add(cell_id)
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def iter_block_paragraphs(container) -> Iterator[Paragraph]:
    element = getattr(container, "element", None)
    if element is None:
        element = container._element
    parent = element.body if element.tag == qn("w:document") else element
    for child in parent:
        if child.tag == qn("w:p"):
            yield Paragraph(child, container)
        elif child.tag == qn("w:tbl"):
            yield from iter_table_paragraphs(Table(child, container))


def iter_header_footer_paragraphs(document: Document) -> Iterator[Paragraph]:
    for section in document.sections:
        for part in (section.header, section.footer):
            if part.is_linked_to_previous:
                continue
            yield from iter_block_paragraphs(part)


def iter_document_paragraphs(document: Document) -> Iterator[Paragraph]:
    yield from iter_block_paragraphs(document)
    yield from iter_header_footer_paragraphs(document)


def collect_document_paragraphs(document: Document) -> list[Paragraph]:
    return list(iter_document_paragraphs(document))


def set_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    if paragraph.text == new_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.add_run(new_text)


class DocxExtractor:
    def __init__(self, config: ExtractorConfig | None = None) -> None:
        self._config = config or DEFAULT_CONFIG.extractor

    def extract(self, docx_path: str | Path) -> ExtractedDocument:
        path = Path(docx_path).expanduser().resolve()
        if not path.is_file():
            raise ExtractionError(f"DOCX not found: {path}")
        if path.suffix.lower() != ".docx":
            raise ExtractionError(f"Expected a .docx file, got: {path.name}")

        logger.info("Opening DOCX: %s", path)
        try:
            document = Document(path)
        except Exception as exc:
            raise ExtractionError(f"Unable to open DOCX {path}: {exc}") from exc

        paragraphs = collect_document_paragraphs(document)
        blocks: list[PageContent] = []
        for index, paragraph in enumerate(paragraphs, start=1):
            text = paragraph.text
            if self._config.repair_pdf_artifacts:
                text = repair_pdf_artifacts(text)
            if not text.strip() and not self._config.keep_empty_pages:
                continue
            blocks.append(PageContent(page_number=index, text=text))

        extracted = ExtractedDocument(source_path=path, pages=tuple(blocks))
        non_empty = sum(1 for block in blocks if block.text.strip())
        logger.info(
            "Extracted %s text block(s) from %s (%s with text, %s characters)",
            extracted.page_count,
            path.name,
            non_empty,
            sum(len(block.text) for block in blocks),
        )
        return extracted


def extract_docx(docx_path: str | Path, config: AppConfig | None = None) -> ExtractedDocument:
    app_config = config or DEFAULT_CONFIG
    return DocxExtractor(app_config.extractor).extract(docx_path)
