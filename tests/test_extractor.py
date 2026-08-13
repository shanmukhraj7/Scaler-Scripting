from pathlib import Path

from docx import Document

from src.exporter import DocxExporter
from src.extractor import extract_docx
from src.faker_utils import FakeValueFactory
from src.models import PIIType, Replacement
from src.replacer import ReplacementEngine
from src.utils import repair_pdf_artifacts


def test_extracts_prospectus_docx():
    path = Path("input/Red Herring Prospectus.docx")
    if not path.exists():
        return
    document = extract_docx(path)
    assert document.page_count > 100
    assert document.pages[0].page_number == 1
    combined = document.full_text
    assert "KSH INTERNATIONAL" in combined
    assert "Sarthak Malvadkar" in combined
    assert "cs.connect@kshinternational.com" in combined


def test_repair_www_and_email():
    text = "www.kshinternational\n. com and user@mail.co\nm"
    fixed = repair_pdf_artifacts(text)
    assert "www.kshinternational.com" in fixed
    assert "user@mail.com" in fixed
    assert "www.kshinternational.com" in repair_pdf_artifacts("www.kshinternational. com")


def test_docx_keeps_table_after_replacement(tmp_path):
    source = tmp_path / "sample.docx"
    output = tmp_path / "redacted.docx"
    doc = Document()
    doc.add_paragraph("Contact Rashi Patil today.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "rashhi.patil@gmail.com"
    table.cell(0, 1).text = "unchanged"
    doc.save(source)

    extracted = extract_docx(source)
    engine = ReplacementEngine(FakeValueFactory())
    engine.mapping.put(
        Replacement("rashhi.patil@gmail.com", "john.doe@example.com", PIIType.EMAIL_ADDRESS)
    )
    engine.mapping.put(Replacement("Rashi Patil", "John Doe", PIIType.PERSON))
    anonymized = engine.anonymize(extracted, [])
    DocxExporter().export(anonymized, output, source_path=source, engine=engine)

    result = Document(output)
    assert len(result.tables) == 1
    assert result.tables[0].cell(0, 1).text.strip() == "unchanged"
    assert result.tables[0].cell(0, 0).text.strip() == "john.doe@example.com"
    assert "John Doe" in result.paragraphs[0].text
    assert "Rashi Patil" not in result.paragraphs[0].text
